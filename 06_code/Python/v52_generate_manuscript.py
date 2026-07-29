#!/usr/bin/env python3
"""
V5.2: Generate BMC Geriatrics Word manuscript
"""
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("Installing python-docx...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

root = r"E:\公共数据库\中国数据库\医养结合政策DID_CHFS_CFPS"

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# ============================================================
# Title
# ============================================================
title = doc.add_heading(
    'Average Policy Effects and State-Dependent Health Mobility During the Expansion of '
    'Integrated Health and Social Care in China: Evidence From Three National Surveys',
    level=1
)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================
# Abstract
# ============================================================
doc.add_heading('Abstract', level=2)

abstract_text = (
    "Background: China's integrated health and social care policy aims to coordinate medical and "
    "care services for older adults. Average policy evaluations may obscure heterogeneity in "
    "health maintenance and recovery across different baseline health states.\n\n"
    "Methods: We used three national surveys with distinct evidence roles. CFPS provided a "
    "pilot-area quasi-experiment estimating the average policy effect on poor self-rated health "
    "using difference-in-differences (DID). CHARLS constructed a 22-item, 6-domain non-disability "
    "health-deficit index for adults aged ≥65 years and compared health-state transition dynamics "
    "before (2011–2015) and during (2015–2018) national policy expansion using discrete-time "
    "multinomial transition analysis with model-standardised common-horizon probabilities. CLASS "
    "provided repeated cross-sectional corroboration of health-deficit associations with concurrent "
    "functional limitations.\n\n"
    "Results: The CFPS pilot-area DID for poor self-rated health was −2.25 percentage points "
    "(95% CI: −10.42 to 5.92; P=0.586), indicating an imprecise average estimate. In CHARLS, "
    "model-standardised two-year transition analysis revealed a state-dependent maintenance–recovery "
    "asymmetry during the expansion period: adults starting in a low-deficit state were 8.7 "
    "percentage points more likely to remain low-deficit (95% bootstrap CI: 4.6 to 13.2), whereas "
    "recovery from intermediate to low-deficit state decreased by 4.0 percentage points (95% CI: "
    "−5.1 to −2.6), and recovery from high to intermediate-deficit state decreased by 6.8 "
    "percentage points (95% CI: −8.2 to −5.5). High-deficit persistence increased by 7.3 "
    "percentage points (95% CI: 5.8 to 8.6). Age-related disadvantage persisted across all "
    "periods without clear widening. CLASS confirmed a strong concurrent association between "
    "health-deficit burden and ADL-help requirement (prevalence ratio per 0.10 FI: 2.06; 95% CI: "
    "1.98 to 2.13).\n\n"
    "Conclusions: The average pilot-area policy estimate was imprecise. Model-standardised CHARLS "
    "analyses suggested state-dependent health-state dynamics during national expansion: improved "
    "maintenance among adults initially in favourable states, but reduced recovery among those in "
    "intermediate or high-deficit states. These findings do not establish a causal national policy "
    "effect but indicate that service expansion may need to distinguish maintaining favourable "
    "health from restoring health among already vulnerable older adults."
)
doc.add_paragraph(abstract_text)

# Keywords
kw = doc.add_paragraph()
kw.add_run('Keywords: ').bold = True
kw.add_run('integrated health and social care; policy evaluation; health-state transitions; '
           'frailty index; difference-in-differences; older adults; China')

# ============================================================
# 1. Introduction
# ============================================================
doc.add_heading('1. Introduction', level=2)

intro_paras = [
    "China's rapidly ageing population, with over 280 million adults aged 60 years or older by "
    "2025, has created unprecedented demand for coordinated medical and long-term care services. "
    "Chronic disease prevalence, functional limitations, and care needs increase substantially "
    "with age, particularly among adults aged 75 years or older [1,2]. In response, China launched "
    "the integrated health and social care policy, beginning with pilot areas in 2016 and expanding "
    "nationally thereafter [3].",

    "Policy evaluations of integrated care typically focus on average outcomes such as self-rated "
    "health, activities of daily living, healthcare utilisation, or mortality [4,5]. However, "
    "average effects may conceal important heterogeneity: the policy may support maintenance of "
    "favourable health states while failing to promote recovery among those already in poor health "
    "[6,7]. Understanding whether health-state dynamics differ by baseline vulnerability is "
    "essential for targeting service delivery and evaluating equity implications.",

    "Traditional static outcome measures cannot distinguish between avoiding deterioration and "
    "promoting recovery. A dynamic framework that tracks individuals across health states can "
    "reveal whether transitions into and out of deficit states changed during policy expansion, "
    "and whether these changes were uniform across age and socioeconomic groups.",

    "This study uses three national surveys with distinct evidence roles to evaluate health "
    "changes during China's integrated care expansion. CFPS provides the core quasi-experimental "
    "policy estimate using a pilot-area difference-in-differences design. CHARLS constructs a "
    "validated health-deficit index and compares health-state transition dynamics before and "
    "during national expansion. CLASS provides repeated cross-sectional corroboration of "
    "health-deficit associations with concurrent functional limitations."
]
for para in intro_paras:
    doc.add_paragraph(para)

# ============================================================
# 2. Methods
# ============================================================
doc.add_heading('2. Methods', level=2)

doc.add_heading('2.1 Study design and evidence architecture', level=3)
doc.add_paragraph(
    "We employed a three-cohort complementary evidence design. Each survey承担 a distinct "
    "analytical role. CFPS served as the core pilot-area quasi-experiment. CHARLS provided "
    "longitudinal health-state transition analysis. CLASS provided repeated cross-sectional "
    "corroboration. The three surveys were not pooled."
)

doc.add_heading('2.2 Policy context', level=3)
doc.add_paragraph(
    "China's integrated health and social care policy was announced in two batches in 2016 "
    "and expanded nationally thereafter. For CFPS, the pre-period comprised 2012 and 2014, "
    "and the post-period comprised 2018. For CHARLS, pre-expansion intervals were 2011–2013 "
    "and 2013–2015, and the national expansion interval was 2015–2018."
)

doc.add_heading('2.3 CFPS data and quasi-experimental analysis', level=3)
doc.add_paragraph(
    "CFPS is a nationally representative household survey. We included adults aged ≥65 years "
    "with at least one pre-period and one post-period observation. The primary outcome was "
    "poor self-rated health. The primary model was individual and survey-year fixed-effects DID."
)

doc.add_heading('2.4 CHARLS data and health-deficit index construction', level=3)
doc.add_paragraph(
    "CHARLS is a nationally representative longitudinal survey. We included adults aged ≥65 "
    "years with at least one observation in 2011, 2013, 2015, or 2018. A 22-item, 6-domain "
    "non-disability health-deficit index was constructed: chronic diseases (9 items), self-rated "
    "health (1 item), depression (1 item), cognition (3 items), physical function (5 items), "
    "and psychiatric conditions (2 items). Items scoring 0 to 1 were summed and divided by the "
    "number of completed eligible items (minimum 80% completion). States were classified as "
    "low-deficit (FI < 0.10), intermediate-deficit (0.10 ≤ FI < 0.25), and high-deficit (FI ≥ 0.25)."
)

doc.add_heading('2.5 CHARLS transition analysis', level=3)
doc.add_paragraph(
    "Discrete-time multinomial health-state transition analysis was used. The model: "
    "next_state ~ current_state × period + interval_years + age_c + female. "
    "Model-standardised two-year transition probabilities were computed for pre-expansion "
    "and expansion-period scenarios. Policy period and interval duration are perfectly "
    "collinear; common-horizon estimates are model-based scenario predictions."
)

doc.add_heading('2.6 Bootstrap analysis', level=3)
doc.add_paragraph(
    "Participant-level cluster bootstrap with 500 successful replications. Participants "
    "resampled with replacement; all transition records retained per participant. Single "
    "multinomial model fitted per replication. Fixed standardisation target. 24 independent "
    "R processes in parallel. All original point estimates fell within 95% bootstrap CI."
)

doc.add_heading('2.7 Sensitivity analyses', level=3)
doc.add_paragraph(
    "History-adjusted transition model; 70% and 90% item-completion thresholds; "
    "sex-stratified analyses; participants with ≥3 and =4 valid waves."
)

doc.add_heading('2.8 CLASS repeated cross-sectional analysis', level=3)
doc.add_paragraph(
    "CLASS could not support longitudinal analysis due to incompatible participant "
    "identifiers across waves. For each available wave, concurrent ADL-help prevalence "
    "by FI level and modified Poisson prevalence ratios were estimated."
)

doc.add_heading('2.9 Statistical software', level=3)
doc.add_paragraph(
    "All analyses were conducted in R (version 4.4.3). Multinomial models used the nnet "
    "package. Bootstrap used independent Rscript processes on 24 logical CPU cores."
)

# ============================================================
# 3. Results
# ============================================================
doc.add_heading('3. Results', level=2)

doc.add_heading('3.1 Sample characteristics', level=3)
doc.add_paragraph(
    "Table 1 presents the study design, sample sizes, and variable architecture for each survey."
)

doc.add_heading('3.2 CFPS primary policy analysis', level=3)
doc.add_paragraph(
    "The primary DID for poor self-rated health was −2.25 percentage points (95% CI: −10.42 "
    "to 5.92; P = 0.586). The average policy estimate was imprecise. The age ≥75 years DDD "
    "was −10.51 percentage points (95% CI: −23.93 to 2.92; conventional P = 0.127; wild-cluster "
    "P = 0.048). The discordance between conventional and wild-cluster inferences indicates "
    "sensitivity to the clustering structure."
)

doc.add_heading('3.3 CHARLS health-deficit index and state validation', level=3)
doc.add_paragraph(
    "The corrected 22-item FI state distribution in 2018 (age ≥65 years): low-deficit "
    "551 (8.3%), intermediate-deficit 2,261 (34.2%), high-deficit 3,802 (57.5%). "
    "State validation against 2018 incident ADL: low-deficit 3.5%, intermediate-deficit "
    "9.2% (RR = 2.66), high-deficit 24.7% (RR = 7.15)."
)

doc.add_heading('3.4 Model-standardised transition results', level=3)
doc.add_paragraph(
    "After common-horizon standardisation to two-year intervals, the principal transition "
    "differences were:"
)
doc.add_paragraph(
    "Low-deficit start: Maintenance of the low-deficit state increased by 8.7 percentage "
    "points (95% bootstrap CI: 4.6 to 13.2). The low-to-intermediate transition decreased "
    "by 7.4 percentage points (95% CI: −11.7 to −3.1).", style='List Bullet'
)
doc.add_paragraph(
    "Intermediate-deficit start: Recovery to low-deficit decreased by 4.0 percentage points "
    "(95% CI: −5.1 to −2.6). Intermediate-state persistence increased by 6.2 percentage "
    "points (95% CI: 3.8 to 8.8).", style='List Bullet'
)
doc.add_paragraph(
    "High-deficit start: Recovery to intermediate-deficit decreased by 6.8 percentage points "
    "(95% CI: −8.2 to −5.5). High-deficit persistence increased by 7.3 percentage points "
    "(95% CI: 5.8 to 8.6).", style='List Bullet'
)

doc.add_heading('3.5 Age and social distributional results', level=3)
doc.add_paragraph(
    "Age ≥75 years was associated with less favourable transitions in both periods. However, "
    "the age-group difference did not clearly widen or narrow during the expansion period. "
    "Existing age-related disadvantage persisted without clear evidence of change."
)

doc.add_heading('3.6 CLASS repeated cross-sectional corroboration', level=3)
doc.add_paragraph(
    "CLASS 2018 (n = 11,163) confirmed a strong concurrent association between health-deficit "
    "burden and ADL-help requirement: prevalence ratio per 0.10 FI = 2.06 (95% CI: 1.98 to 2.13)."
)

# ============================================================
# 4. Discussion
# ============================================================
doc.add_heading('4. Discussion', level=2)

doc.add_paragraph(
    "This three-survey study evaluated health changes during China's integrated health and "
    "social care expansion. The CFPS pilot-area average policy estimate was imprecise. "
    "CHARLS model-standardised analyses revealed a state-dependent maintenance–recovery "
    "asymmetry during the expansion period: adults starting in a low-deficit state were more "
    "likely to remain there, whereas recovery from intermediate- and high-deficit states was "
    "less frequent. Age-related disadvantage persisted across all periods without clear "
    "evidence of widening. CLASS confirmed a strong concurrent association between "
    "health-deficit burden and ADL-help requirement."
)

doc.add_paragraph(
    "The findings suggest that the expansion of integrated care services may have been more "
    "effective at supporting health maintenance among adults in favourable states than at "
    "promoting recovery among those already in intermediate or high-deficit states. These "
    "results do not imply that the policy was harmful or ineffective. Rather, they indicate "
    "that average policy evaluations may obscure important heterogeneity in how different "
    "population subgroups experience health changes during service expansion."
)

doc.add_paragraph(
    "The CFPS DID provides an average pilot-area estimate that is imprecise. The CHARLS "
    "transition analysis reveals that this average may conceal opposing dynamics: improved "
    "maintenance among low-deficit adults coexisting with reduced recovery among "
    "higher-deficit adults."
)

doc.add_paragraph(
    "Strengths of this study include the use of three national surveys with distinct evidence "
    "roles, a validated multi-domain health-deficit index, complete state-transition modelling, "
    "500-replication participant-level cluster bootstrap, and systematic sensitivity analyses "
    "including history-dependence testing."
)

doc.add_paragraph(
    "Limitations include: (1) CFPS pilot-area matching relies on limited pre-trend information; "
    "(2) CHARLS lacks an untreated national comparison group; (3) policy period and interval "
    "duration are perfectly collinear; (4) the first-order Markov assumption is imperfect; "
    "(5) death and attrition may introduce selection; (6) CLASS cannot support longitudinal "
    "analysis; (7) CHARLS and CLASS use different FI item sets."
)

doc.add_paragraph(
    "The average pilot-area policy estimate was imprecise. Model-standardised CHARLS analyses "
    "suggested a state-dependent maintenance–recovery asymmetry during the national expansion "
    "period: adults starting in a low-deficit state were more likely to remain there, whereas "
    "recovery from intermediate- and high-deficit states was less frequent. These findings do "
    "not establish a causal national policy effect but indicate that service expansion may need "
    "to distinguish maintaining favourable health from restoring health among already vulnerable "
    "older adults."
)

# ============================================================
# 5. Code Availability
# ============================================================
doc.add_heading('5. Code Availability', level=2)
doc.add_paragraph(
    'The statistical analysis code used for data processing, variable construction, model '
    'estimation, sensitivity analyses, and figure generation is publicly available at '
    'https://github.com/r-ruser/china-policy-did.git.'
)

# ============================================================
# Tables
# ============================================================
doc.add_page_break()
doc.add_heading('Tables', level=2)

# Table 1
doc.add_heading('Table 1. Study design, sample characteristics, and variable architecture', level=3)
table1 = doc.add_table(rows=6, cols=4)
table1.style = 'Table Grid'
table1.cell(0, 0).text = 'Feature'
table1.cell(0, 1).text = 'CFPS'
table1.cell(0, 2).text = 'CHARLS'
table1.cell(0, 3).text = 'CLASS'
table1.cell(1, 0).text = 'Design'
table1.cell(1, 1).text = 'Pilot-area quasi-experiment'
table1.cell(1, 2).text = 'Longitudinal cohort'
table1.cell(1, 3).text = 'Repeated cross-sectional'
table1.cell(2, 0).text = 'Role'
table1.cell(2, 1).text = 'Core policy evidence'
table1.cell(2, 2).text = 'Transition dynamics'
table1.cell(2, 3).text = 'External corroboration'
table1.cell(3, 0).text = 'Sample'
table1.cell(3, 1).text = 'Adults ≥65, pilot vs non-pilot'
table1.cell(3, 2).text = 'Adults ≥65, 4 waves'
table1.cell(3, 3).text = 'Adults ≥65, 4 waves'
table1.cell(4, 0).text = 'Primary outcome'
table1.cell(4, 1).text = 'Poor self-rated health'
table1.cell(4, 2).text = 'Health-deficit state'
table1.cell(4, 3).text = 'ADL-help requirement'
table1.cell(5, 0).text = 'Analysis'
table1.cell(5, 1).text = 'DID/DDD'
table1.cell(5, 2).text = 'Multinomial transition'
table1.cell(5, 3).text = 'Modified Poisson'

# Table 2
doc.add_heading('Table 2. CFPS policy estimates', level=3)
table2 = doc.add_table(rows=4, cols=4)
table2.style = 'Table Grid'
table2.cell(0, 0).text = 'Estimand'
table2.cell(0, 1).text = 'Estimate (pp)'
table2.cell(0, 2).text = '95% CI'
table2.cell(0, 3).text = 'P'
table2.cell(1, 0).text = 'Pilot-area DID'
table2.cell(1, 1).text = '−2.25'
table2.cell(1, 2).text = '−10.42 to 5.92'
table2.cell(1, 3).text = '0.586'
table2.cell(2, 0).text = 'Age ≥75 DDD'
table2.cell(2, 1).text = '−10.51'
table2.cell(2, 2).text = '−23.93 to 2.92'
table2.cell(2, 3).text = '0.127'
table2.cell(3, 0).text = 'Social vulnerability DDD'
table2.cell(3, 1).text = '[value]'
table2.cell(3, 2).text = '[CI]'
table2.cell(3, 3).text = '[P]'

# Table 3
doc.add_heading('Table 3. CHARLS state distribution and ADL validation', level=3)
table3 = doc.add_table(rows=4, cols=5)
table3.style = 'Table Grid'
table3.cell(0, 0).text = 'State'
table3.cell(0, 1).text = '2018 N'
table3.cell(0, 2).text = '2018 %'
table3.cell(0, 3).text = '2018 ADL risk'
table3.cell(0, 4).text = 'RR'
table3.cell(1, 0).text = 'Low-deficit'
table3.cell(1, 1).text = '551'
table3.cell(1, 2).text = '8.3'
table3.cell(1, 3).text = '3.5%'
table3.cell(1, 4).text = '1.00'
table3.cell(2, 0).text = 'Intermediate'
table3.cell(2, 1).text = '2,261'
table3.cell(2, 2).text = '34.2'
table3.cell(2, 3).text = '9.2%'
table3.cell(2, 4).text = '2.66'
table3.cell(3, 0).text = 'High-deficit'
table3.cell(3, 1).text = '3,802'
table3.cell(3, 2).text = '57.5'
table3.cell(3, 3).text = '24.7%'
table3.cell(3, 4).text = '7.15'

# Table 4
doc.add_heading('Table 4. CHARLS common-horizon transition probability differences', level=3)
table4 = doc.add_table(rows=10, cols=4)
table4.style = 'Table Grid'
table4.cell(0, 0).text = 'Start'
table4.cell(0, 1).text = 'Transition'
table4.cell(0, 2).text = 'Difference (pp)'
table4.cell(0, 3).text = '95% Bootstrap CI'
transitions = [
    ('Low', 'Maintain Low', '+8.7', '4.6 to 13.2'),
    ('Low', 'Low → Intermediate', '−7.4', '−11.7 to −3.1'),
    ('Low', 'Low → High', '−1.4', '−3.8 to 0.9'),
    ('Intermediate', 'Intermediate → Low', '−4.0', '−5.1 to −2.6'),
    ('Intermediate', 'Maintain Intermediate', '+6.2', '3.8 to 8.8'),
    ('Intermediate', 'Intermediate → High', '−2.3', '−4.6 to 0.0'),
    ('High', 'High → Low', '−0.5', '−0.7 to −0.2'),
    ('High', 'High → Intermediate', '−6.8', '−8.2 to −5.5'),
    ('High', 'Maintain High', '+7.3', '5.8 to 8.6'),
]
for i, (start, trans, diff, ci) in enumerate(transitions):
    table4.cell(i+1, 0).text = start
    table4.cell(i+1, 1).text = trans
    table4.cell(i+1, 2).text = diff
    table4.cell(i+1, 3).text = ci

# ============================================================
# Figure Legends
# ============================================================
doc.add_page_break()
doc.add_heading('Figure Legends', level=2)

doc.add_paragraph(
    'Figure 1. Study design and evidence architecture. Three national surveys with distinct '
    'roles: CFPS (core pilot-area quasi-experiment), CHARLS (longitudinal health-state '
    'transition analysis), CLASS (repeated cross-sectional corroboration).'
)

doc.add_paragraph(
    'Figure 2. CFPS policy estimates. (A) Pilot-area DID and age ≥75 DDD for poor '
    'self-rated health. (B) Event-study diagnostic. (C) Robustness analyses.'
)

doc.add_paragraph(
    'Figure 3. CHARLS common-horizon transition probability differences. Model-standardised '
    'two-year probability differences (expansion minus pre-expansion) for each start state. '
    'Error bars show 95% bootstrap confidence intervals from 500 replications.'
)

# ============================================================
# Save
# ============================================================
output_path = os.path.join(root, "V5.2_full_manuscript_BMC_Geriatrics.docx")
doc.save(output_path)
print(f"Saved: {output_path}")

# Generate clean version (same content, no track changes)
doc2 = Document(output_path)
output_clean = os.path.join(root, "V5.2_clean_manuscript_without_track_changes.docx")
doc2.save(output_clean)
print(f"Saved: {output_clean}")

print("Manuscript generation complete.")
