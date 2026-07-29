#!/usr/bin/env python3
"""V5.2: Generate BMC Geriatrics Methods & Results only (double-spaced)"""
import os
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

root = r"E:\公共数据库\中国数据库\医养结合政策DID_CHFS_CFPS"
doc = Document()

# Set double-spaced, Times New Roman 12pt
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
pf = style.paragraph_format
pf.line_spacing = 2.0
pf.space_after = Pt(0)

# Helper to add double-spaced paragraph
def add_para(text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if bold:
        run.bold = True
    return p

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.line_spacing = 2.0
    for run in h.runs:
        run.font.name = 'Times New Roman'
    return h

# ============================================================
# Title
# ============================================================
title = doc.add_heading(
    'Average Policy Effects and State-Dependent Health Mobility During the '
    'Expansion of Integrated Health and Social Care in China: Evidence From '
    'Three National Surveys', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

# ============================================================
# Methods
# ============================================================
add_heading('Methods', level=2)

add_para('2.1 Study design and evidence architecture', bold=True)
add_para(
    'We employed a three-cohort complementary evidence design. CFPS served as the core '
    'pilot-area quasi-experiment estimating average policy effects. CHARLS provided '
    'longitudinal health-state transition analysis before and during national policy '
    'expansion. CLASS provided repeated cross-sectional corroboration of health-deficit '
    'associations with concurrent functional limitations. The three surveys承担 distinct '
    'analytical roles and were not pooled.')

add_para('2.2 Policy context', bold=True)
add_para(
    'China\'s integrated health and social care policy was announced in two batches in 2016 '
    'and expanded nationally thereafter. For CFPS, the pre-period comprised 2012 and 2014, '
    'and the post-period comprised 2018. For CHARLS, pre-expansion intervals were 2011–2013 '
    'and 2013–2015, and the national expansion interval was 2015–2018. CLASS waves '
    '(2016, 2018, 2020, 2023) all occurred during or after national expansion.')

add_para('2.3 CFPS data and quasi-experimental analysis', bold=True)
add_para(
    'CFPS is a nationally representative household survey covering 25 provinces. We included '
    'adults aged ≥65 years with at least one pre-period (2012 or 2014) and one post-period '
    '(2018) observation. The primary outcome was poor self-rated health (self-rated health '
    '≥4 on a 5-point scale). The primary model was individual and survey-year fixed-effects '
    'difference-in-differences (DID):')
add_para(
    '    poor_srh_it = α_i + γ_t + β(PilotCity_i × Post_t) + ε_it',
    bold=False)
add_para(
    'Distributional models used triple-difference (DDD) with baseline age ≥75 years and '
    'current social vulnerability as effect modifiers. Robustness checks included '
    'city-clustered standard errors, wild-cluster bootstrap (999 replicates), ATT weighting, '
    'leave-one-city-out, leave-one-province-out, and equivalence analyses.')

add_para('2.4 CHARLS data and health-deficit index construction', bold=True)
add_para(
    'CHARLS is a nationally representative longitudinal survey covering 28 provinces. We '
    'included adults aged ≥65 years with at least one observation in 2011, 2013, 2015, or '
    '2018. A 22-item, 6-domain non-disability health-deficit index was constructed: chronic '
    'diseases (9 items: hypertension, heart disease, stroke, lung disease, diabetes, cancer, '
    'arthritis, kidney disease, liver disease), self-rated health (1 item), depression '
    '(CESD-10 total score, 1 domain-level composite), cognition (3 items: orientation, '
    'immediate word recall, serial-7 subtraction), physical function (5 items: stooping, '
    'walking 1 km, lifting, standing from chair, climbing stairs), and psychiatric conditions '
    '(2 items). ADL/IADL items were excluded as external outcomes. Current smoking was excluded '
    'as a behavioural exposure. Each deficit was scored 0 (no deficit) to 1 (deficit present). '
    'The FI was computed as the sum of observed deficit scores divided by the number of completed '
    'eligible items, requiring ≥80% item completion. Values ranged from 0 to 1.')

add_para(
    'States were classified as: low-deficit (FI < 0.10), intermediate-deficit (0.10 ≤ FI < 0.25), '
    'and high-deficit (FI ≥ 0.25). Cut-points were prespecified based on deficit-accumulation '
    'conventions and validated against subsequent ADL limitation. Neutral state labels were used '
    'throughout.')

add_para('2.5 CHARLS state validation', bold=True)
add_para(
    'States were validated against subsequent ADL limitation. Among 2015 participants without '
    'baseline ADL limitation, the 2018 incident ADL risk was 3.5% for low-deficit, 9.2% '
    '(RR = 2.66; 95% CI: 2.10 to 3.36) for intermediate-deficit, and 24.7% (RR = 7.15; '
    '95% CI: 5.69 to 8.94) for high-deficit states.')

add_para('2.6 CHARLS transition analysis', bold=True)
add_para(
    'Discrete-time multinomial health-state transition analysis was used. For each '
    'participant-interval, the next observed state was modelled conditional on the current '
    'state, policy period, and covariates:')
add_para(
    '    next_state ~ current_state × period + interval_years + age_c + female')
add_para(
    'Model-standardised two-year transition probabilities were computed for both '
    'pre-expansion and expansion-period scenarios, with interval_years fixed at 2 and '
    'covariates set to reference values (age 70, 50% female). Policy period and interval '
    'duration are perfectly collinear in the observed data; common-horizon estimates are '
    'model-based scenario predictions, not independently adjusted policy-period effects.')

add_para('2.7 Bootstrap analysis', bold=True)
add_para(
    'Participant-level cluster bootstrap with 500 successful replications was performed '
    'using 24 independent R processes on all available logical CPU cores. Participants '
    'were resampled with replacement; all transition records were retained per participant. '
    'A single multinomial model (identical to the primary analysis) was fitted per '
    'replication with a fixed standardisation target. The identity test confirmed original '
    'estimates were reproduced within 10⁻⁶. All original point estimates fell within '
    'the 95% bootstrap percentile intervals, with maximum bias of 0.0011.')

add_para('2.8 Sensitivity analyses', bold=True)
add_para(
    'A history-adjusted transition model included the previous observed state as a covariate. '
    'Additional sensitivity analyses used 70% and 90% item-completion thresholds, '
    'sex-stratified analyses, and participants with at least three or exactly four valid waves.')

add_para('2.9 CLASS repeated cross-sectional analysis', bold=True)
add_para(
    'CLASS could not support longitudinal analysis due to incompatible participant '
    'identifiers across waves. For each available wave, concurrent ADL-help prevalence '
    'by FI level and modified Poisson prevalence ratios were estimated. Childhood hunger '
    '(b14) was analysed as a cross-sectional association where available.')

add_para('2.10 Statistical software', bold=True)
add_para(
    'All analyses were conducted in R (version 4.4.3). Multinomial models used the nnet '
    'package. Bootstrap used independent Rscript processes on 24 logical CPU cores.')

# ============================================================
# Results
# ============================================================
add_heading('Results', level=2)

add_para('3.1 Sample characteristics and evidence architecture', bold=True)
add_para(
    'Table 1 presents the study design, sample sizes, and variable architecture for each survey.')

# Table 1
t1 = doc.add_table(rows=6, cols=4)
t1.style = 'Table Grid'
headers = ['Feature', 'CFPS', 'CHARLS', 'CLASS']
for i, h in enumerate(headers):
    t1.cell(0, i).text = h
rows_data = [
    ['Design', 'Pilot-area quasi-experiment', 'Longitudinal cohort', 'Repeated cross-sectional'],
    ['Role', 'Core policy evidence', 'Transition dynamics', 'External corroboration'],
    ['Sample', 'Adults ≥65, pilot vs non-pilot', 'Adults ≥65, 4 waves', 'Adults ≥65, 4 waves'],
    ['Primary outcome', 'Poor self-rated health', 'Health-deficit state', 'ADL-help requirement'],
    ['Analysis', 'DID/DDD', 'Multinomial transition', 'Modified Poisson'],
]
for r, row in enumerate(rows_data):
    for c, val in enumerate(row):
        t1.cell(r+1, c).text = val

add_para('3.2 CFPS primary policy analysis', bold=True)
add_para(
    'The primary DID for poor self-rated health was −2.25 percentage points (95% CI: '
    '−10.42 to 5.92; P = 0.586; Table 2). The 90% equivalence interval did not lie '
    'entirely within ±5 percentage points, and the 80% minimum detectable effect was '
    '7.68 percentage points. The average policy estimate was therefore imprecise.')
add_para(
    'The age ≥75 years DDD was −10.51 percentage points (95% CI: −23.93 to 2.92; '
    'conventional P = 0.127; wild-cluster P = 0.048). The discordance between conventional '
    'and wild-cluster inferences indicates sensitivity to the clustering structure.')

# Table 2
t2 = doc.add_table(rows=4, cols=4)
t2.style = 'Table Grid'
for i, h in enumerate(['Estimand', 'Estimate (pp)', '95% CI', 'P']):
    t2.cell(0, i).text = h
t2.cell(1, 0).text = 'Pilot-area DID'
t2.cell(1, 1).text = '−2.25'
t2.cell(1, 2).text = '−10.42 to 5.92'
t2.cell(1, 3).text = '0.586'
t2.cell(2, 0).text = 'Age ≥75 DDD'
t2.cell(2, 1).text = '−10.51'
t2.cell(2, 2).text = '−23.93 to 2.92'
t2.cell(2, 3).text = '0.127'
t2.cell(3, 0).text = 'Social vulnerability DDD'
t2.cell(3, 1).text = '[value]'
t2.cell(3, 2).text = '[CI]'
t2.cell(3, 3).text = '[P]'

add_para('3.3 CHARLS health-deficit index and state validation', bold=True)
add_para(
    'The corrected 22-item FI had the following state distribution in 2018 (age ≥65 years, '
    '80% completion threshold): low-deficit 551 (8.3%), intermediate-deficit 2,261 (34.2%), '
    'high-deficit 3,802 (57.5%). State classification was validated against 2018 incident ADL '
    'among 2015 participants without baseline ADL: low-deficit 3.5%, intermediate-deficit '
    '9.2% (RR = 2.66), high-deficit 24.7% (RR = 7.15; Table 3).')

# Table 3
t3 = doc.add_table(rows=4, cols=5)
t3.style = 'Table Grid'
for i, h in enumerate(['State', '2018 N', '2018 %', 'ADL risk', 'RR']):
    t3.cell(0, i).text = h
t3.cell(1, 0).text = 'Low-deficit'
t3.cell(1, 1).text = '551'
t3.cell(1, 2).text = '8.3'
t3.cell(1, 3).text = '3.5%'
t3.cell(1, 4).text = '1.00'
t3.cell(2, 0).text = 'Intermediate'
t3.cell(2, 1).text = '2,261'
t3.cell(2, 2).text = '34.2'
t3.cell(2, 3).text = '9.2%'
t3.cell(2, 4).text = '2.66'
t3.cell(3, 0).text = 'High-deficit'
t3.cell(3, 1).text = '3,802'
t3.cell(3, 2).text = '57.5'
t3.cell(3, 3).text = '24.7%'
t3.cell(3, 4).text = '7.15'

add_para('3.4 Model-standardised transition results', bold=True)
add_para(
    'After common-horizon standardisation to two-year intervals, the principal transition '
    'differences were (Table 4, Figure 3):')
add_para(
    'Low-deficit start: Maintenance of the low-deficit state increased by 8.7 percentage '
    'points (95% bootstrap CI: 4.6 to 13.2). The low-to-intermediate transition decreased '
    'by 7.4 percentage points (95% CI: −11.7 to −3.1). The low-to-high transition '
    'showed a small non-significant decrease (−1.4 pp; 95% CI: −3.8 to 0.9).')
add_para(
    'Intermediate-deficit start: Recovery to low-deficit decreased by 4.0 percentage points '
    '(95% CI: −5.1 to −2.6). Intermediate-state persistence increased by 6.2 '
    'percentage points (95% CI: 3.8 to 8.8). The intermediate-to-high transition showed '
    'a small non-significant decrease (−2.3 pp; 95% CI: −4.6 to 0.0).')
add_para(
    'High-deficit start: Recovery to intermediate-deficit decreased by 6.8 percentage '
    'points (95% CI: −8.2 to −5.5). High-deficit persistence increased by 7.3 '
    'percentage points (95% CI: 5.8 to 8.6). The high-to-low transition was very small '
    'and non-significant (−0.5 pp; 95% CI: −0.7 to −0.2).')
add_para(
    'The overall pattern indicates a state-dependent maintenance–recovery asymmetry: '
    'improved maintenance among adults initially in favourable states, but reduced recovery '
    'among those in intermediate or high-deficit states.')

# Table 4
t4 = doc.add_table(rows=10, cols=4)
t4.style = 'Table Grid'
for i, h in enumerate(['Start', 'Transition', 'Difference (pp)', '95% Bootstrap CI']):
    t4.cell(0, i).text = h
transitions = [
    ('Low', 'Maintain Low', '+8.7', '4.6 to 13.2'),
    ('Low', 'Low → Intermediate', '−7.4', '−11.7 to −3.1'),
    ('Low', 'Low → High', '−1.4', '−3.8 to 0.9'),
    ('Intermediate', 'Intermediate → Low', '−4.0', '−5.1 to −2.6'),
    ('Intermediate', 'Maintain Intermediate', '+6.2', '3.8 to 8.8'),
    ('Intermediate', 'Intermediate → High', '−2.3', '−4.6 to 0.0'),
    ('High', 'High → Low', '−0.5', '−0.7 to −0.2'),
    ('High', 'High → Intermediate', '−6.8', '−8.2 to −5.5'),
    ('High', 'Maintain High', '+7.3', '5.8 to 8.6'),
]
for i, (start, trans, diff, ci) in enumerate(transitions):
    t4.cell(i+1, 0).text = start
    t4.cell(i+1, 1).text = trans
    t4.cell(i+1, 2).text = diff
    t4.cell(i+1, 3).text = ci

add_para(
    'All values are model-standardised common-horizon two-year probability differences '
    '(expansion period minus pre-expansion). 95% CIs from 500-replication participant-level '
    'cluster bootstrap.', bold=False)

# Figure 3 placeholder
doc.add_page_break()
add_para('Figure 3. CHARLS common-horizon transition probability differences.', bold=True)
add_para(
    'Model-standardised two-year probability differences (expansion minus pre-expansion) '
    'for each start state. Error bars show 95% bootstrap confidence intervals from '
    '500 replications. Panel A: Low-deficit start. Panel B: Intermediate-deficit start. '
    'Panel C: High-deficit start.')

add_para('3.5 Age and social distributional results', bold=True)
add_para(
    'Age ≥75 years was associated with less favourable transitions in both periods. '
    'However, the age-group difference did not clearly widen or narrow during the expansion '
    'period. Existing age-related disadvantage persisted without clear evidence of change.')

add_para('3.6 History-adjusted and sensitivity analyses', bold=True)
add_para(
    'The history-adjusted model (including previous state as a covariate) had substantially '
    'better fit than the first-order model (AIC: 7,358 vs 7,829 on identical observations). '
    'The main transition patterns persisted in the history-adjusted analysis, though some '
    'magnitudes attenuated. All sensitivity analyses (70%/90% completion thresholds, '
    'sex-stratified, ≥3-wave and =4-wave subsamples) showed the same directional pattern.')

add_para('3.7 CLASS repeated cross-sectional corroboration', bold=True)
add_para(
    'CLASS 2018 (n = 11,163) confirmed a strong concurrent association between health-deficit '
    'burden and ADL-help requirement: prevalence ratio per 0.10 FI = 2.06 (95% CI: 1.98 to '
    '2.13). This association persisted after excluding physical-function items from the FI. '
    'Adults aged ≥75 years had higher FI levels and greater ADL-help prevalence than '
    'those aged 65–74 years.')

# ============================================================
# Code Availability
# ============================================================
add_heading('Code Availability', level=2)
add_para(
    'The statistical analysis code used for data processing, variable construction, model '
    'estimation, sensitivity analyses, and figure generation is publicly available at '
    'https://github.com/r-ruser/china-policy-did.git.')

# Save
out = os.path.join(root, "V5.2_full_manuscript_BMC_Geriatrics_methods_results.docx")
doc.save(out)
print(f"Saved: {out}")
